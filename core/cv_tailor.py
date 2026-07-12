"""
cv_tailor.py — Tailors a CV for a specific job WITHOUT inventing anything.

Only runs on jobs that already scored >= match_threshold_percent (see
matcher.py). Its job is narrow and honest:
  1. Pick the best-fitting summary variant from skill_bank.json
  2. Reorder skills so JD-relevant ones appear first
  3. Reorder projects/experience so the most relevant ones appear first
  4. Lightly rephrase bullets to mirror the JD's own wording — WITHOUT
     changing what was actually done (this is what closes the gap from an
     80-90% match to a much stronger keyword/phrasing alignment — it is not
     a guarantee of literally 100%, and it will never claim skills you don't have)
  5. Render the result as a clean .docx

This keeps a hard rule: nothing in the output CV can name a skill, tool, or
achievement that isn't already present somewhere in skill_bank.json.
"""

import os
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.llm_client import generate_text


def _all_allowed_terms(skill_bank):
    """Every real term the candidate is allowed to be described with."""
    terms = {s["name"].lower() for s in skill_bank.get("skills", [])}
    return terms


def rewrite_bullets_with_ai(bullets, job_description, allowed_terms, api_key=None):
    """
    Ask the AI to rephrase bullets to mirror the JD's language, with an
    explicit instruction not to add any new tools/skills/claims.
    Falls back to the original bullets untouched if every LLM provider fails.
    """
    prompt = f"""Rewrite these resume bullet points so their PHRASING mirrors
the language and keywords used in the job description below. Do NOT add any
new skill, tool, technology, number, or achievement that isn't already in the
original bullets. Only rephrase — same facts, wording closer to the JD.

JOB DESCRIPTION:
{job_description[:2000]}

ORIGINAL BULLETS:
{chr(10).join('- ' + b for b in bullets)}

Return ONLY the rewritten bullets, one per line, no numbering, no extra text."""

    try:
        raw_text = generate_text(prompt, gemini_key=api_key)
        new_bullets = [line.strip("- ").strip() for line in raw_text.strip().split("\n") if line.strip()]

        # Safety check: reject the rewrite if it introduces a term not in the
        # allowed skill list AND not present in any original bullet (crude but effective).
        original_text = " ".join(bullets).lower()
        for nb in new_bullets:
            words = set(re.findall(r"[a-zA-Z][a-zA-Z0-9+.#]{2,}", nb.lower()))
            new_terms = words - allowed_terms
            # allow common English words through; only flag if a *tech-looking* new term appears
            suspicious = [w for w in new_terms if w not in original_text and len(w) > 3]
            if len(suspicious) > 2:  # a couple of stray words is normal English, not fabrication
                print(f"[cv_tailor] Rewrite looked suspicious, keeping original bullet: {nb[:50]}...")
                return bullets  # bail out to the safe original set

        return new_bullets if len(new_bullets) == len(bullets) else bullets

    except Exception as e:
        print(f"[cv_tailor] Bullet rewrite failed ({e}), keeping original wording")
        return bullets


def pick_summary_variant(skill_bank, matching_skills):
    """Pick whichever summary variant best matches the job's matched skills."""
    variants = skill_bank.get("summary_variants", {})
    matched_lower = " ".join(matching_skills).lower()

    if any(term in matched_lower for term in ["testing", "qa", "quality assurance", "postman", "unit test", "sqa"]):
        return variants.get("qa", variants.get("general", ""))
    if any(term in matched_lower for term in ["react", "css", "html", "frontend", "ui", "next.js"]):
        return variants.get("frontend", variants.get("general", ""))
    if any(term in matched_lower for term in ["node", "sql", "api", "backend", "server", "c#", ".net"]):
        return variants.get("backend", variants.get("general", ""))
    return variants.get("general", "")


def build_tailored_cv_data(job_title, job_description, skill_bank, match_result, api_key=None):
    """
    Returns a plain dict describing the tailored CV content — reordered and
    reworded, but every fact traceable back to skill_bank.json.
    """
    allowed_terms = _all_allowed_terms(skill_bank)
    matching_skills = match_result.get("matching_skills", [])

    # Reorder skills: matched-relevant ones first, rest after
    all_skills = [s["name"] for s in skill_bank.get("skills", [])]
    matched_set = {m.lower() for m in matching_skills}
    ordered_skills = (
        [s for s in all_skills if s.lower() in matched_set] +
        [s for s in all_skills if s.lower() not in matched_set]
    )

    # Reorder + rewrite projects: relevance = word-level overlap between matched
    # skills and the project's tags/bullets — exact-phrase matching was too
    # strict (e.g. "Unit Testing (NUnit / xUnit basics)" as a skill name rarely
    # appears verbatim in a bullet even when the underlying content matches).
    import re as _re

    def _significant_words(text):
        words = _re.findall(r"[a-zA-Z][a-zA-Z0-9#+.]{2,}", text.lower())
        stopwords = {"the", "and", "for", "with", "using", "via", "basics"}
        return {w for w in words if w not in stopwords}

    matching_words = set()
    for skill in matching_skills:
        matching_words |= _significant_words(skill)

    def relevance(item):
        text = " ".join(item.get("tags", [])) + " " + " ".join(item.get("bullets", []))
        item_words = _significant_words(text)
        return len(matching_words & item_words)

    projects = sorted(skill_bank.get("projects", []), key=relevance, reverse=True)
    experience = sorted(skill_bank.get("experience", []), key=relevance, reverse=True)

    # Cap to the most relevant few — a junior/intern CV should be ~1 page.
    # Including all projects every time regardless of relevance produced a
    # 3-page CV in testing, which reads as unfocused rather than thorough.
    MAX_PROJECTS = 4
    MAX_EXPERIENCE = 3
    projects = projects[:MAX_PROJECTS]
    experience = experience[:MAX_EXPERIENCE]

    for p in projects:
        p = dict(p)
        p["bullets"] = rewrite_bullets_with_ai(p["bullets"], job_description, allowed_terms, api_key)
    for e in experience:
        e = dict(e)
        e["bullets"] = rewrite_bullets_with_ai(e["bullets"], job_description, allowed_terms, api_key)

    return {
        "summary": pick_summary_variant(skill_bank, matching_skills),
        "skills": ordered_skills,
        "projects": projects,
        "experience": experience,
        "education": skill_bank.get("education", {}),
        "certifications": skill_bank.get("certifications", []),
        "job_title": job_title,
    }


def render_cv_docx(candidate, tailored_data, output_path):
    """Render the tailored CV data into a clean, professional-looking .docx file."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ACCENT_COLOR = RGBColor(0x1F, 0x3A, 0x5F)  # dark navy — professional, ATS-safe
    BODY_FONT = "Calibri"

    doc = Document()

    # Tighter margins so more real content fits on one page
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Set a sane default font for the whole document instead of Word's default Calibri-but-inconsistent-sizing
    normal_style = doc.styles["Normal"]
    normal_style.font.name = BODY_FONT
    normal_style.font.size = Pt(10.5)
    normal_style.paragraph_format.space_after = Pt(4)

    def add_section_heading(text):
        """A heading with a bottom border line under it, like a real resume section divider."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.font.bold = True
        run.font.size = Pt(11.5)
        run.font.color.rgb = ACCENT_COLOR
        run.font.name = BODY_FONT

        p_border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F3A5F")
        p_border.append(bottom)
        p._p.get_or_add_pPr().append(p_border)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.name = BODY_FONT
        return p

    # ── Header: name + contact line ──
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(candidate.get("name", "Your Name"))
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = ACCENT_COLOR
    name_run.font.name = BODY_FONT
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_bits = [candidate.get("email", ""), candidate.get("phone", ""), candidate.get("location", "")]
    links = candidate.get("links", {})
    contact_bits += [v for v in links.values() if v]
    contact_p = doc.add_paragraph(" | ".join(b for b in contact_bits if b))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(8)
    for run in contact_p.runs:
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Summary ──
    add_section_heading("Summary")
    summary_p = doc.add_paragraph(tailored_data["summary"])
    for run in summary_p.runs:
        run.font.size = Pt(10.5)

    # ── Skills ──
    add_section_heading("Skills")
    skills_p = doc.add_paragraph(" • ".join(tailored_data["skills"]))
    for run in skills_p.runs:
        run.font.size = Pt(10)

    # ── Projects ──
    if tailored_data["projects"]:
        add_section_heading("Projects")
        for p in tailored_data["projects"]:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(6)
            title_p.paragraph_format.space_after = Pt(2)
            title_run = title_p.add_run(p["title"])
            title_run.bold = True
            title_run.font.size = Pt(10.5)
            title_run.font.name = BODY_FONT
            for bullet in p["bullets"]:
                add_bullet(bullet)

    # ── Experience ──
    if tailored_data["experience"]:
        add_section_heading("Experience")
        for e in tailored_data["experience"]:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(6)
            title_p.paragraph_format.space_after = Pt(2)
            title_run = title_p.add_run(e["title"])
            title_run.bold = True
            title_run.font.size = Pt(10.5)
            title_run.font.name = BODY_FONT
            for bullet in e["bullets"]:
                add_bullet(bullet)

    # ── Education ──
    edu = tailored_data.get("education", {})
    if edu:
        add_section_heading("Education")
        edu_p1 = doc.add_paragraph(f"{edu.get('degree','')} — {edu.get('institute','')}")
        edu_p1.runs[0].bold = True
        edu_p1.runs[0].font.size = Pt(10.5)
        edu_p2 = doc.add_paragraph(f"Graduating {edu.get('graduation_year','')}  |  CGPA {edu.get('cgpa','')}")
        edu_p2.runs[0].font.size = Pt(10)
        if edu.get("coursework"):
            edu_p3 = doc.add_paragraph(f"Relevant coursework: {', '.join(edu['coursework'])}")
            edu_p3.runs[0].font.size = Pt(9.5)
            edu_p3.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Certifications ──
    certifications = tailored_data.get("certifications", [])
    if certifications:
        add_section_heading("Certifications")
        for cert in certifications:
            add_bullet(cert)

    doc.save(output_path)
    print(f"[cv_tailor] Tailored CV saved -> {output_path}")
    return output_path


if __name__ == "__main__":
    from dotenv import load_dotenv
    load_dotenv()

    with open("data/skill_bank.json") as f:
        bank = json.load(f)

    candidate = {"name": "Test Candidate", "email": "test@example.com", "location": "Remote"}
    fake_match = {"matching_skills": ["React", "HTML / CSS", "Git / GitHub"], "score": 85}

    jd = "Looking for a Junior Frontend Developer intern who knows React and CSS. Remote, 6-person team."
    tailored = build_tailored_cv_data("Junior Frontend Developer Intern", jd, bank, fake_match)
    render_cv_docx(candidate, tailored, "test_tailored_cv.docx")
